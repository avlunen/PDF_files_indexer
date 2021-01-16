#-----------------------------------------------------------------------
# Name: Index Printer
# Purpose: To read a pdf file and find a specified index list, recording
# the page number within the PDF where the terms occur
#
# Author: Alexander von Lunen
#
# Created: 24/07/2019
# Copyright: (c) Alexander von Lunen 2019
# Licence: Public Domain
#-----------------------------------------------------------------------
import glob
import os
import time
import io
import re
import sys
import PyPDF2
import collections
import pdftotext
import string
import unicodedata
from docx import Document
#from string import maketrans

tbl = dict.fromkeys(i for i in xrange(sys.maxunicode)
                      if unicodedata.category(unichr(i)).startswith('P'))

# strip punctuation from unicode strings
def remove_punctuation(text):
    return text.translate(tbl)

# read TSV file with index terms
def init_index_list(ind_list):
    with io.open('all_concordances.txt', mode='r', encoding='utf-8') as tsv:
        indl = [line.strip().split('\t') for line in tsv]

        for i in indl:
            if len(i) == 2:
                ind_list[i[0]] = i[1]
            elif len(i) == 1:
                ind_list[i[0]] = i[0]
            else:
                continue

# get page offset, i.e. we don't want to index the front matter, such as the
# preface; for this, we check the PDF page labels, and only start when there
# are arabic numberals, as roman numerals indicate the front matter
def get_pdf_offset(pdf):
    # get page label array, pairs of physical page number and IndirectObjects
    page_labels = pdf.trailer["/Root"]["/PageLabels"]["/Nums"]
    for i in xrange(0, len(page_labels), 2): # iterate through the pairs
        if page_labels[i+1].getObject()['/S'] == '/D': # if an arabic numeral is indicated
            return page_labels[i]

    return 0

# export index
def export_index(ind_list):
    document = Document()

    document.add_heading('Index', 0)

    f = ''

    for i, e in ind_list.iteritems():
        first_letter = i[0]
        if first_letter.upper() != f.upper():
            document.add_heading(first_letter.upper(), level=1)
        p = document.add_paragraph(i + ' ' + ', '.join(str(pe) for pe in e))
        f = first_letter

    document.save('index.docx')

# main routine
def main():
    index_list = {}
    page_list = {}
    pdf_file = open('proofs2.pdf', 'rb') # read the PDF you want to search, will need customization
    read_pdf = PyPDF2.PdfFileReader(pdf_file)
    page_offset = get_pdf_offset(read_pdf)

    # pyPDF2 has trouble reading CMap chars, use pdftotext
    with open('proofs2.pdf', 'rb') as f: # read the PDF you want to search, will need customization
        pdf2 = pdftotext.PDF(f)

    # get number of pages
    number_of_pages = read_pdf.getNumPages()
    print "Number of pages: %d" % number_of_pages

    # read list of index terms from external file
    print "Reading index terms ..."
    init_index_list(index_list)

    # create list of unified index terms by removing duplicates from values
    pl = {k: [] for k in list(set(index_list.values()))}
    page_list = collections.OrderedDict(sorted(pl.items(), key=lambda t: t[0].lower()))
    del pl

    print "Reading pages ..."
    # go through PDF page by page
    for x in range(number_of_pages):
        if x <= page_offset: # skip intro etc. May needs changing
            continue
        page = read_pdf.getPage(x)
        page_number = read_pdf.getPageNumber(page)
        page_label = page_number-page_offset+1
        print "\tPhysical page: %d, Page label: %d" % (page_number,page_label)

        # get content of the page, remove punctutation so that the search terms get found
        page_content = ' '.join(remove_punctuation(pdf2[page_number]).split())

        # check each entry in the index list against the page's content
        for i, e in index_list.iteritems():
             if i in page_content: # found index term, add page number to its entry
                #print i
                if page_label not in page_list[e]: # avoid duplicates
                     page_list[e].append(page_label)

    print "Printing index ..."
    # output index
    export_index(page_list)

    # uncomment for debugging
    #for i, e in page_list.iteritems():
        #print i.encode('utf-8'), e

    print "... Done!"

if __name__ == '__main__':
   main()
